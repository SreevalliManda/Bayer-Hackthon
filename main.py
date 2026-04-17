import json
import os
from io import BytesIO
from fastapi import FastAPI
from fastapi.responses import StreamingResponse
from app.graph import incident_graph, IncidentState
from dotenv import load_dotenv
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

load_dotenv()

app = FastAPI(
    title="Incident AI - Advanced Multi-Agent System",
    description="AI-powered incident investigation system using LangChain + OpenAI agents",
    version="2.0",
    docs_url="/docs",
    redoc_url="/redoc"
)

def load_data():
    """Load mock data from JSON files"""
    with open("app/data/metrics.json") as f:
        metrics_data = json.load(f)
    
    with open("app/data/logs.json") as f:
        logs_data = json.load(f)
    
    with open("app/data/deploy.json") as f:
        deploy_data = json.load(f)
    
    with open("app/data/alerts.json") as f:
        alerts_data = json.load(f)
    
    return metrics_data, logs_data, deploy_data, alerts_data

@app.get(
    "/incident",
    tags=["Incident Analysis"],
    summary="Trigger full incident analysis",
    description="""
    Analyze incident data across multiple dimensions:
    - **Metrics**: Latency spikes (checkout service latency spike to 2000ms at 10:15)
    - **Logs**: DB connection timeouts correlated with the latency spike
    - **Deployments**: Configuration change 15 minutes before spike (DB config update at 10:00)
    - **Alerts**: Critical alerts for db_timeout, latency_spike, error_rate, service_down
    
    Returns comprehensive investigation findings with root cause analysis and recommendations.
    
    **Example Scenario**:
    - Issue: Checkout service latency spikes to 2000ms
    - Root Cause: DB config deployment (10:00) changed connection pool settings
    - Effect: DB connection timeouts cascade at 10:14-10:16, causing latency spike at 10:15
    - Recommendation: Immediate rollback to previous configuration
    """
)
async def handle_incident():
    try:
        # Load data
        metrics_data, logs_data, deploy_data, alerts_data = load_data()
        
        # Create initial state
        initial_state = IncidentState(
            metrics_data=metrics_data,
            logs_data=logs_data,
            deploy_data=deploy_data,
            alerts_data=alerts_data
        )
        
        # Run the incident analysis graph
        final_state = incident_graph.invoke(initial_state)
        
        # Convert dict to IncidentState if necessary
        if isinstance(final_state, dict):
            final_state = IncidentState(**final_state)
        
        # Return comprehensive results
        return {
            "status": "success",
            "investigation_plan": final_state.investigation_plan,
            "analysis": {
                "metrics_analysis": final_state.metrics_analysis,
                "logs_analysis": final_state.logs_analysis,
                "deploy_analysis": final_state.deploy_analysis,
                "correlation": final_state.correlation,
                "decision": final_state.decision,
                "confidence": f"{final_state.confidence * 100:.1f}%"
            },
            "report": final_state.report,
            "raw_data": {
                "metrics_count": len(metrics_data),
                "logs_count": len(logs_data),
                "deployments_count": len(deploy_data),
                "alerts_count": len(alerts_data)
            }
        }
        
    except Exception as e:
        return {
            "status": "error",
            "message": str(e),
            "details": "Ensure OPENAI_API_KEY is set in .env file"
        }

@app.get(
    "/health",
    tags=["System"],
    summary="Health check endpoint"
)
async def health_check():
    """System health check - indicates if the service is running"""
    return {"status": "healthy", "service": "Incident AI Advanced"}

@app.get(
    "/incident/report/download",
    tags=["Incident Analysis"],
    summary="Download investigation report as Word document",
    description="""
    Generate and download a professionally formatted incident report in DOCX format.
    
    The report includes:
    - **Investigation Plan**: Alert assessment and investigation priorities
    - **Metrics Analysis**: Performance anomalies and degradation patterns
    - **Logs Analysis**: Error patterns and stack traces
    - **Deployment Analysis**: Correlation between deployments and errors
    - **Commander Decision**: Root cause determination and confidence level
    - **Final Report**: Executive summary with recommendations
    
    **Scenario Detection**:
    - Identifies latency spike (2000ms) in checkout service
    - Correlates with DB config deployment 15 minutes prior
    - Recommends immediate rollback
    - Confidence: High (based on temporal correlation)
    """
)
async def download_incident_report():
    try:
        # Load data
        metrics_data, logs_data, deploy_data, alerts_data = load_data()
        
        # Create initial state
        initial_state = IncidentState(
            metrics_data=metrics_data,
            logs_data=logs_data,
            deploy_data=deploy_data,
            alerts_data=alerts_data
        )
        
        # Run the incident analysis graph
        final_state = incident_graph.invoke(initial_state)
        
        # Convert dict to IncidentState if necessary
        if isinstance(final_state, dict):
            final_state = IncidentState(**final_state)
        
        # Create Word document
        doc = Document()
        
        # Add title
        title = doc.add_heading('Incident Analysis Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add investigation plan section
        doc.add_heading('Investigation Plan', level=1)
        doc.add_paragraph(final_state.investigation_plan)
        
        # Add analysis sections
        doc.add_heading('Detailed Analysis', level=1)
        
        doc.add_heading('Metrics Analysis', level=2)
        doc.add_paragraph(final_state.metrics_analysis)
        
        doc.add_heading('Logs Analysis', level=2)
        doc.add_paragraph(final_state.logs_analysis)
        
        doc.add_heading('Deployment Analysis', level=2)
        doc.add_paragraph(final_state.deploy_analysis)
        
        # Add correlation and decision
        doc.add_heading('Commander Decision', level=1)
        doc.add_paragraph(final_state.decision)
        
        confidence_para = doc.add_paragraph()
        confidence_run = confidence_para.add_run(f"Confidence Level: {final_state.confidence * 100:.1f}%")
        confidence_run.bold = True
        
        # Add final report
        doc.add_heading('Final Report', level=1)
        doc.add_paragraph(final_state.report)
        
        # Add metadata footer
        doc.add_paragraph()
        footer = doc.add_paragraph("---")
        footer_text = doc.add_paragraph(f"Generated by Incident AI | OpenAI-powered Multi-Agent System")
        footer_text.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Save to BytesIO
        output = BytesIO()
        doc.save(output)
        output.seek(0)
        
        # Return as file download using StreamingResponse
        return StreamingResponse(
            iter([output.getvalue()]),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": "attachment; filename=incident_report.docx"}
        )
        
    except Exception as e:
        return {
            "status": "error",
            "message": str(e),
            "details": "Ensure OPENAI_API_KEY is set in .env file"
        }

@app.get(
    "/",
    tags=["System"],
    summary="API documentation and endpoints"
)
async def root():
    """Root endpoint - displays available endpoints and system information"""
    return {
        "message": "Incident AI - Advanced Multi-Agent Incident Response System",
        "version": "2.0",
        "endpoints": {
            "/incident": "Trigger full incident analysis",
            "/incident/report/download": "Generate report and download as Word document",
            "/health": "Service health check"
        },
        "powered_by": "LangGraph + OpenAI"
    }